import re
# Import docx NOT python-docx
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
import datetime

def geraListaCapitulos(last):
    data = [str(i) for i in range(1,last+1)]
    return data

def leCapitulos(caps_file_name):
    f = open(caps_file_name, "r", encoding="utf-8")
    lines = f.readlines()

    for i, title in enumerate(lines):
        lines[i] = title.replace("\n", '')
    
    return lines

def geraDocx(file_name, book_title, author, weekendMessage, caps_file_name, skipWeekend, day, month, year):
    # Create an instance of a word document
    doc = docx.Document()

    #Centered
    title = doc.add_paragraph()
    paragraph_format = title.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add a Title to the document
    title = title.add_run(book_title + '\n' + author, 0)

    #Font size
    title.font.size = Pt(20)
    
    # Table data in a form of list
    #data = geraListaCapitulos(28)
    data = leCapitulos(caps_file_name)
    
    # Creating a table object
    table = doc.add_table(rows=1, cols=3)

    DATE = 0
    CAP = 1
    STATUS = 2
    
    # Adding heading in the 1st row of the table
    row = table.rows[0].cells
    row[DATE].text = 'Data'
    row[CAP].text = 'Meta de leitura'
    row[STATUS].text = 'Status'

    #Heading bold
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Adding data from the list to the table
    NONE = ""
    ROW_HEIGHT = 0.7

    current_date = datetime.datetime(year, month, day) - datetime.timedelta(days=1)

    for cap in data:
        #update day
        current_date = current_date + datetime.timedelta(days=1)

        #skip weekend
        while(skipWeekend and current_date.weekday() >= 5):
            # Adding a row and then adding data in it.
            row = table.add_row().cells

            row[DATE].text = current_date.strftime("%d/%m/%y") 
            row[CAP].text = weekendMessage
            row[STATUS].text = NONE

            table.rows[-1].height = Cm(ROW_HEIGHT)

            current_date = current_date + datetime.timedelta(days=1)

        # Adding a row and then adding data in it.
        row = table.add_row().cells

        # Converting id to string as table can only take string input
        row[DATE].text = current_date.strftime("%d/%m/%y") 
        row[CAP].text = cap
        row[STATUS].text = NONE
 
        table.rows[-1].height = Cm(ROW_HEIGHT)

    #Style to a table
    table.style = 'Table Grid'
    
    # Now save the document to a location
    doc.save(f"output/{file_name}.docx")
